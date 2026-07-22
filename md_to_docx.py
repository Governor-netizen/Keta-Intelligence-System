#!/usr/bin/env python3
"""
Convert the Keta manuscript markdown into an editable Word document.
Handles: ATX headings, pipe tables, ![](img) embeds, **bold**, *italic*,
numbered/bulleted lists, horizontal rules, and bold "Figure N."/"Table N."
caption paragraphs. Plain-text equation lines are kept as monospace.
"""
import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE = os.path.dirname(os.path.abspath(__file__))
SRC = os.path.join(BASE, "keta_manuscript_v4_2_draft.md")
OUT = os.path.join(BASE, "keta_manuscript_v4_2.docx")

INLINE = re.compile(r"(\*\*.+?\*\*|\*.+?\*|_.+?_)")


def add_runs(paragraph, text):
    """Add text to a paragraph, honoring **bold** and *italic*."""
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("*") and part.endswith("*"):
            paragraph.add_run(part[1:-1]).italic = True
        elif part.startswith("_") and part.endswith("_"):
            paragraph.add_run(part[1:-1]).italic = True
        else:
            paragraph.add_run(part)


def is_table_row(line):
    return line.strip().startswith("|") and line.strip().endswith("|")


def parse_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    with open(SRC, encoding="utf-8") as f:
        lines = f.read().split("\n")

    i = 0
    title_done = False
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # blank
        if not stripped:
            i += 1
            continue

        # horizontal rule
        if stripped == "---":
            i += 1
            continue

        # image ![alt](path)
        m = re.match(r"!\[.*?\]\((.+?)\)", stripped)
        if m:
            img = m.group(1)
            path = img if os.path.isabs(img) else os.path.join(BASE, img)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if os.path.exists(path):
                try:
                    p.add_run().add_picture(path, width=Inches(5.8))
                except Exception:
                    p.add_run(f"[image: {img}]").italic = True
            else:
                p.add_run(f"[missing image: {img}]").italic = True
            i += 1
            continue

        # headings
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped.lstrip("#").strip()
            if level == 1 and not title_done:
                h = doc.add_heading("", level=0)
                add_runs(h, text)
                title_done = True
            else:
                h = doc.add_heading("", level=min(level, 4))
                add_runs(h, text)
            i += 1
            continue

        # tables
        if is_table_row(line) and i + 1 < len(lines) and is_table_row(lines[i + 1]) \
                and set(lines[i + 1].strip()) <= set("|-: "):
            header = parse_row(line)
            rows = []
            i += 2  # skip header + separator
            while i < len(lines) and is_table_row(lines[i]):
                rows.append(parse_row(lines[i]))
                i += 1
            ncol = len(header)
            table = doc.add_table(rows=1, cols=ncol)
            table.style = "Light Grid Accent 1"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for c, cell_text in enumerate(header):
                cell = table.rows[0].cells[c]
                cell.paragraphs[0].text = ""
                run = cell.paragraphs[0].add_run(cell_text.replace("**", ""))
                run.bold = True
            for r in rows:
                cells = table.add_row().cells
                for c in range(ncol):
                    val = r[c] if c < len(r) else ""
                    cells[c].paragraphs[0].text = ""
                    add_runs(cells[c].paragraphs[0], val)
            doc.add_paragraph()
            continue

        # list items
        m = re.match(r"^(\d+)\.\s+(.*)", stripped)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_runs(p, m.group(2))
            i += 1
            continue
        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, stripped[2:])
            i += 1
            continue

        # equation-ish lines (contain math tokens, no prose punctuation run)
        if re.search(r"[=^]|sqrt|sigma|omega|mu_|I_diff|TWL\(t\)", stripped) \
                and len(stripped) < 90 and not stripped.endswith("."):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stripped)
            run.font.name = "Consolas"
            run.font.size = Pt(10.5)
            i += 1
            continue

        # figure/table caption -> keep bold lead, smaller
        if stripped.startswith("**Figure") or stripped.startswith("**Table"):
            p = doc.add_paragraph()
            add_runs(p, stripped)
            for r in p.runs:
                r.font.size = Pt(9.5)
            i += 1
            continue

        # default paragraph
        p = doc.add_paragraph()
        add_runs(p, stripped)
        i += 1

    doc.save(OUT)
    print(f"Saved {OUT}")


if __name__ == "__main__":
    main()
